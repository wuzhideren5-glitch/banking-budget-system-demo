from __future__ import annotations

import json
import re
import shutil
import ast
import operator
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import aiosqlite
from docx import Document
from fastapi import HTTPException, UploadFile

from app.db_paths import budget_db_path, common_db_path
from app.schemas import (
    SmartReportCalcMetricComponent,
    SmartReportCalcMetricRow,
    SmartReportCalcMetricUpsert,
    SmartReportGenerateRequest,
    SmartReportGenerateResponse,
    SmartReportInstanceRow,
    SmartReportPreviewRequest,
    SmartReportPreviewResponse,
    SmartReportTemplateCreateResponse,
    SmartReportTemplateRow,
    SmartReportTemplateVariableRow,
    SmartReportTemplateVariableUpsert,
)


PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def _iso_now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _safe_code(raw: str) -> str:
    code = re.sub(r"[^A-Za-z0-9_-]+", "_", raw.strip())
    code = re.sub(r"_+", "_", code).strip("_")
    if not code:
        raise HTTPException(status_code=400, detail="模板编码不能为空")
    return code[:80]


def _json_loads_dict(raw: str | None) -> dict[str, Any]:
    if not raw:
        return {}
    try:
        value = json.loads(raw)
    except Exception:
        return {}
    return value if isinstance(value, dict) else {}


def _month_label(value: Any) -> str | None:
    if value is None or str(value).strip() == "":
        return None
    text = str(value).strip().upper()
    if re.fullmatch(r"M\d{2}", text):
        return text
    if re.fullmatch(r"\d{1,2}", text):
        num = int(text)
        if 1 <= num <= 12:
            return f"M{num:02d}"
    return text


def _year_label(value: Any) -> str:
    text = str(value or "").strip().upper()
    if re.fullmatch(r"Y\d{4}", text):
        return text
    if re.fullmatch(r"\d{4}", text):
        return f"Y{text}"
    return f"Y{datetime.now().year}"


def _plain_year(value: Any) -> int:
    label = _year_label(value)
    m = re.search(r"(\d{4})", label)
    return int(m.group(1)) if m else datetime.now().year


def _format_number(value: float, *, percentage: bool = False) -> str:
    if percentage:
        return f"{value * 100:.2f}%"
    abs_value = abs(value)
    if abs_value >= 10000:
        return f"{value:,.2f}"
    if value == int(value):
        return f"{int(value):,}"
    return f"{value:,.2f}"


_CALC_OPERATORS: dict[type[ast.AST], Any] = {
    ast.Add: operator.add,
    ast.Sub: operator.sub,
    ast.Mult: operator.mul,
    ast.Div: operator.truediv,
    ast.Pow: operator.pow,
    ast.USub: operator.neg,
    ast.UAdd: operator.pos,
}


class SmartReportService:
    def __init__(self, *, data_dir: Path) -> None:
        self.data_dir = data_dir
        self.template_dir = data_dir / "smart_report_templates"
        self.output_dir = data_dir / "smart_report_outputs"

    async def list_templates(self) -> list[SmartReportTemplateRow]:
        async with aiosqlite.connect(common_db_path()) as db:
            await db.execute("PRAGMA foreign_keys = ON")
            cur = await db.execute(
                """
                SELECT t.template_id, t.template_code, t.template_name, t.template_type,
                       t.status, t.version_no, t.remark, t.created_at, t.updated_at,
                       COUNT(v.variable_id) AS variable_count
                FROM smart_report_template t
                LEFT JOIN smart_report_template_variable v ON v.template_id = t.template_id
                GROUP BY t.template_id
                ORDER BY t.updated_at DESC, t.template_id DESC
                """
            )
            rows = await cur.fetchall()
        return [
            SmartReportTemplateRow(
                template_id=int(r[0]),
                template_code=str(r[1]),
                template_name=str(r[2]),
                template_type=str(r[3]),
                status=str(r[4]),
                version_no=int(r[5]),
                remark=str(r[6]) if r[6] is not None else None,
                created_at=str(r[7]),
                updated_at=str(r[8]),
                variable_count=int(r[9] or 0),
            )
            for r in rows
        ]

    async def create_or_update_template(
        self,
        *,
        file: UploadFile,
        template_code: str,
        template_name: str,
        template_type: str = "analysis",
        remark: str | None = None,
        created_by: str | None = None,
    ) -> SmartReportTemplateCreateResponse:
        if not file.filename or not file.filename.lower().endswith(".docx"):
            raise HTTPException(status_code=400, detail="请上传 .docx Word 模板")
        safe_code = _safe_code(template_code)
        now = _iso_now()
        self.template_dir.mkdir(parents=True, exist_ok=True)

        async with aiosqlite.connect(common_db_path()) as db:
            await db.execute("PRAGMA foreign_keys = ON")
            cur = await db.execute(
                "SELECT template_id, version_no FROM smart_report_template WHERE template_code = ?",
                (safe_code,),
            )
            existing = await cur.fetchone()
            next_version = int(existing[1]) + 1 if existing else 1
            target = self.template_dir / f"{safe_code}_v{next_version}.docx"
            with target.open("wb") as fh:
                shutil.copyfileobj(file.file, fh)

            placeholders = self.extract_placeholders(target)
            if existing:
                template_id = int(existing[0])
                await db.execute(
                    """
                    UPDATE smart_report_template
                    SET template_name = ?, template_type = ?, file_path = ?, version_no = ?,
                        remark = ?, updated_at = ?
                    WHERE template_id = ?
                    """,
                    (template_name.strip(), template_type.strip() or "analysis", str(target), next_version, remark, now, template_id),
                )
            else:
                cur = await db.execute(
                    """
                    INSERT INTO smart_report_template (
                      template_code, template_name, template_type, file_path, status,
                      version_no, remark, created_by, created_at, updated_at
                    ) VALUES (?, ?, ?, ?, 'active', ?, ?, ?, ?, ?)
                    """,
                    (
                        safe_code,
                        template_name.strip() or safe_code,
                        template_type.strip() or "analysis",
                        str(target),
                        next_version,
                        remark,
                        created_by,
                        now,
                        now,
                    ),
                )
                template_id = int(cur.lastrowid)

            await self._sync_detected_variables(db, template_id, placeholders, now)
            await db.commit()

        template = await self.get_template(template_id)
        return SmartReportTemplateCreateResponse(template=template, placeholders=placeholders)

    async def create_or_update_text_template(
        self,
        *,
        template_code: str,
        template_name: str,
        content: str,
        template_type: str = "analysis",
        remark: str | None = None,
        created_by: str | None = None,
    ) -> SmartReportTemplateCreateResponse:
        safe_code = _safe_code(template_code)
        now = _iso_now()
        self.template_dir.mkdir(parents=True, exist_ok=True)

        async with aiosqlite.connect(common_db_path()) as db:
            await db.execute("PRAGMA foreign_keys = ON")
            cur = await db.execute(
                "SELECT template_id, version_no FROM smart_report_template WHERE template_code = ?",
                (safe_code,),
            )
            existing = await cur.fetchone()
            next_version = int(existing[1]) + 1 if existing else 1
            target = self.template_dir / f"{safe_code}_v{next_version}.docx"
            self._write_text_template_docx(target, template_name.strip() or safe_code, content)
            placeholders = self.extract_placeholders(target)

            if existing:
                template_id = int(existing[0])
                await db.execute(
                    """
                    UPDATE smart_report_template
                    SET template_name = ?, template_type = ?, file_path = ?, version_no = ?,
                        remark = ?, updated_at = ?
                    WHERE template_id = ?
                    """,
                    (template_name.strip(), template_type.strip() or "analysis", str(target), next_version, remark, now, template_id),
                )
            else:
                cur = await db.execute(
                    """
                    INSERT INTO smart_report_template (
                      template_code, template_name, template_type, file_path, status,
                      version_no, remark, created_by, created_at, updated_at
                    ) VALUES (?, ?, ?, ?, 'active', ?, ?, ?, ?, ?)
                    """,
                    (
                        safe_code,
                        template_name.strip() or safe_code,
                        template_type.strip() or "analysis",
                        str(target),
                        next_version,
                        remark,
                        created_by,
                        now,
                        now,
                    ),
                )
                template_id = int(cur.lastrowid)

            await self._sync_detected_variables(db, template_id, placeholders, now)
            await db.commit()

        template = await self.get_template(template_id)
        return SmartReportTemplateCreateResponse(template=template, placeholders=placeholders)

    async def get_template(self, template_id: int) -> SmartReportTemplateRow:
        rows = [t for t in await self.list_templates() if t.template_id == template_id]
        if not rows:
            raise HTTPException(status_code=404, detail="报告模板不存在")
        return rows[0]

    async def list_variables(self, template_id: int) -> list[SmartReportTemplateVariableRow]:
        async with aiosqlite.connect(common_db_path()) as db:
            await db.execute("PRAGMA foreign_keys = ON")
            cur = await db.execute(
                """
                SELECT variable_id, template_id, variable_key, variable_name, variable_type,
                       binding_config_json, display_order, created_at, updated_at
                FROM smart_report_template_variable
                WHERE template_id = ?
                ORDER BY display_order, variable_id
                """,
                (template_id,),
            )
            rows = await cur.fetchall()
        return [
            SmartReportTemplateVariableRow(
                variable_id=int(r[0]),
                template_id=int(r[1]),
                variable_key=str(r[2]),
                variable_name=str(r[3]),
                variable_type=str(r[4]),  # type: ignore[arg-type]
                binding_config=_json_loads_dict(r[5]),
                display_order=int(r[6] or 0),
                created_at=str(r[7]),
                updated_at=str(r[8]),
            )
            for r in rows
        ]

    async def upsert_variables(
        self, template_id: int, variables: list[SmartReportTemplateVariableUpsert]
    ) -> list[SmartReportTemplateVariableRow]:
        await self.get_template(template_id)
        now = _iso_now()
        async with aiosqlite.connect(common_db_path()) as db:
            await db.execute("PRAGMA foreign_keys = ON")
            for item in variables:
                inferred_type, inferred_config = self._infer_variable(item.variable_key)
                variable_type = item.variable_type or inferred_type
                binding_config = item.binding_config or inferred_config
                variable_name = item.variable_name or self._default_variable_name(item.variable_key)
                await db.execute(
                    """
                    INSERT INTO smart_report_template_variable (
                      template_id, variable_key, variable_name, variable_type,
                      binding_config_json, display_order, created_at, updated_at
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    ON CONFLICT(template_id, variable_key) DO UPDATE SET
                      variable_name = excluded.variable_name,
                      variable_type = excluded.variable_type,
                      binding_config_json = excluded.binding_config_json,
                      display_order = excluded.display_order,
                      updated_at = excluded.updated_at
                    """,
                    (
                        template_id,
                        item.variable_key.strip(),
                        variable_name,
                        variable_type,
                        json.dumps(binding_config, ensure_ascii=False),
                        item.display_order,
                        now,
                        now,
                    ),
                )
            await db.commit()
        return await self.list_variables(template_id)

    async def list_calc_metrics(self) -> list[SmartReportCalcMetricRow]:
        async with aiosqlite.connect(common_db_path()) as db:
            cur = await db.execute(
                """
                SELECT metric_code, metric_name, expression, components_json, value_type,
                       format_type, remark, created_at, updated_at
                FROM smart_report_calc_metric
                ORDER BY updated_at DESC, metric_code
                """
            )
            rows = await cur.fetchall()
        return [self._calc_metric_from_row(row) for row in rows]

    async def upsert_calc_metric(self, body: SmartReportCalcMetricUpsert) -> SmartReportCalcMetricRow:
        metric_code = re.sub(r"[^A-Za-z0-9_]+", "_", body.metric_code.strip()).strip("_")
        if not metric_code:
            raise HTTPException(status_code=400, detail="计算指标编码不能为空")
        aliases = {item.alias.strip() for item in body.components}
        if not aliases:
            raise HTTPException(status_code=400, detail="请至少选择一个基础指标")
        self._safe_eval_expression(body.expression, {alias: 1.0 for alias in aliases})
        now = _iso_now()
        components = [item.model_dump() for item in body.components]
        async with aiosqlite.connect(common_db_path()) as db:
            await db.execute(
                """
                INSERT INTO smart_report_calc_metric (
                  metric_code, metric_name, expression, components_json, value_type,
                  format_type, remark, created_at, updated_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(metric_code) DO UPDATE SET
                  metric_name = excluded.metric_name,
                  expression = excluded.expression,
                  components_json = excluded.components_json,
                  value_type = excluded.value_type,
                  format_type = excluded.format_type,
                  remark = excluded.remark,
                  updated_at = excluded.updated_at
                """,
                (
                    metric_code,
                    body.metric_name.strip(),
                    body.expression.strip(),
                    json.dumps(components, ensure_ascii=False),
                    body.value_type.strip() or "金额",
                    body.format_type.strip() or "number",
                    body.remark,
                    now,
                    now,
                ),
            )
            await db.commit()
        metric = await self.get_calc_metric(metric_code)
        if not metric:
            raise HTTPException(status_code=500, detail="计算指标保存失败")
        return metric

    async def get_calc_metric(self, metric_code: str) -> SmartReportCalcMetricRow | None:
        async with aiosqlite.connect(common_db_path()) as db:
            cur = await db.execute(
                """
                SELECT metric_code, metric_name, expression, components_json, value_type,
                       format_type, remark, created_at, updated_at
                FROM smart_report_calc_metric
                WHERE metric_code = ?
                """,
                (metric_code.strip(),),
            )
            row = await cur.fetchone()
        return self._calc_metric_from_row(row) if row else None

    async def generate(self, body: SmartReportGenerateRequest) -> SmartReportGenerateResponse:
        now = _iso_now()
        template_row = await self._fetch_template_record(body.template_id)
        instance_name = body.instance_name or f"{template_row['template_name']} {now}"
        self.output_dir.mkdir(parents=True, exist_ok=True)

        async with aiosqlite.connect(common_db_path()) as db:
            await db.execute("PRAGMA foreign_keys = ON")
            cur = await db.execute(
                """
                INSERT INTO smart_report_instance (
                  report_id, template_id, instance_name, parameter_values_json,
                  text_values_json, generation_status, created_at, updated_at
                ) VALUES (?, ?, ?, ?, ?, 'running', ?, ?)
                """,
                (
                    body.report_id,
                    body.template_id,
                    instance_name,
                    json.dumps(body.parameters, ensure_ascii=False),
                    json.dumps(body.text_values, ensure_ascii=False),
                    now,
                    now,
                ),
            )
            instance_id = int(cur.lastrowid)
            cur = await db.execute(
                """
                INSERT INTO smart_report_job (instance_id, job_type, job_status, started_at)
                VALUES (?, 'generate', 'running', ?)
                """,
                (instance_id, now),
            )
            job_id = int(cur.lastrowid)
            await db.commit()

        return await self._render_instance(
            instance_id=instance_id,
            job_id=job_id,
            template_id=body.template_id,
            parameters=body.parameters,
            text_values=body.text_values,
            job_type="generate",
        )

    async def preview(self, body: SmartReportPreviewRequest) -> SmartReportPreviewResponse:
        template_row = await self._fetch_template_record(body.template_id)
        template_path = Path(str(template_row["file_path"]))
        variables = await self._variables_for_template(body.template_id, template_path)
        resolved, warnings = await self._resolve_values(variables, body.parameters, body.text_values)
        return SmartReportPreviewResponse(
            preview_text=self._render_preview_text(template_path, resolved),
            resolved_values=resolved,
            warnings=warnings,
        )

    async def refresh_instance(self, instance_id: int) -> SmartReportGenerateResponse:
        now = _iso_now()
        async with aiosqlite.connect(common_db_path()) as db:
            await db.execute("PRAGMA foreign_keys = ON")
            cur = await db.execute(
                """
                SELECT template_id, parameter_values_json, text_values_json
                FROM smart_report_instance
                WHERE instance_id = ?
                """,
                (instance_id,),
            )
            row = await cur.fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="报告实例不存在")
            template_id = int(row[0])
            parameters = _json_loads_dict(row[1])
            text_values = _json_loads_dict(row[2])
            await db.execute(
                """
                UPDATE smart_report_instance
                SET generation_status = 'running', error_message = NULL, updated_at = ?
                WHERE instance_id = ?
                """,
                (now, instance_id),
            )
            cur = await db.execute(
                """
                INSERT INTO smart_report_job (instance_id, job_type, job_status, started_at)
                VALUES (?, 'refresh', 'running', ?)
                """,
                (instance_id, now),
            )
            job_id = int(cur.lastrowid)
            await db.commit()

        return await self._render_instance(
            instance_id=instance_id,
            job_id=job_id,
            template_id=template_id,
            parameters=parameters,
            text_values=text_values,
            job_type="refresh",
        )

    async def _render_instance(
        self,
        *,
        instance_id: int,
        job_id: int,
        template_id: int,
        parameters: dict[str, Any],
        text_values: dict[str, Any],
        job_type: str,
    ) -> SmartReportGenerateResponse:
        try:
            self.output_dir.mkdir(parents=True, exist_ok=True)
            template_row = await self._fetch_template_record(template_id)
            template_path = Path(str(template_row["file_path"]))
            variables = await self._variables_for_template(template_id, template_path)
            resolved, warnings = await self._resolve_values(variables, parameters, text_values)
            prefix = "smart_report_refresh" if job_type == "refresh" else "smart_report"
            output_filename = f"{prefix}_{instance_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = self.output_dir / output_filename
            self._render_docx(template_path, output_path, resolved)
            finished = _iso_now()
            async with aiosqlite.connect(common_db_path()) as db:
                await db.execute("PRAGMA foreign_keys = ON")
                if job_type == "refresh":
                    await db.execute(
                        """
                        UPDATE smart_report_instance
                        SET output_file_path = ?, data_snapshot_json = ?, generation_status = 'success',
                            last_refresh_at = ?, updated_at = ?
                        WHERE instance_id = ?
                        """,
                        (
                            str(output_path),
                            json.dumps(resolved, ensure_ascii=False),
                            finished,
                            finished,
                            instance_id,
                        ),
                    )
                else:
                    await db.execute(
                        """
                        UPDATE smart_report_instance
                        SET output_file_path = ?, data_snapshot_json = ?, generation_status = 'success',
                            last_generated_at = ?, updated_at = ?
                        WHERE instance_id = ?
                        """,
                        (
                            str(output_path),
                            json.dumps(resolved, ensure_ascii=False),
                            finished,
                            finished,
                            instance_id,
                        ),
                    )
                await db.execute(
                    """
                    UPDATE smart_report_job
                    SET job_status = 'success', finished_at = ?
                    WHERE job_id = ?
                    """,
                    (finished, job_id),
                )
                await db.commit()
            return SmartReportGenerateResponse(
                instance_id=instance_id,
                job_id=job_id,
                output_filename=output_filename,
                download_url=f"/api/smart-reports/instances/{instance_id}/download",
                generated_at=finished,
                resolved_values=resolved,
                warnings=warnings,
            )
        except Exception as exc:
            finished = _iso_now()
            async with aiosqlite.connect(common_db_path()) as db:
                await db.execute("PRAGMA foreign_keys = ON")
                await db.execute(
                    """
                    UPDATE smart_report_instance
                    SET generation_status = 'failed', error_message = ?, updated_at = ?
                    WHERE instance_id = ?
                    """,
                    (str(exc), finished, instance_id),
                )
                await db.execute(
                    """
                    UPDATE smart_report_job
                    SET job_status = 'failed', finished_at = ?, error_message = ?
                    WHERE job_id = ?
                    """,
                    (finished, str(exc), job_id),
                )
                await db.commit()
            if isinstance(exc, HTTPException):
                raise exc
            raise HTTPException(status_code=500, detail=f"生成 Word 报告失败：{exc}") from exc

    async def list_instances(self) -> list[SmartReportInstanceRow]:
        async with aiosqlite.connect(common_db_path()) as db:
            await db.execute("PRAGMA foreign_keys = ON")
            cur = await db.execute(
                """
                SELECT i.instance_id, i.report_id, i.template_id, t.template_name,
                       i.instance_name, i.generation_status, i.output_file_path,
                       i.error_message, i.last_generated_at, i.last_refresh_at,
                       i.created_at, i.updated_at
                FROM smart_report_instance i
                LEFT JOIN smart_report_template t ON t.template_id = i.template_id
                ORDER BY i.updated_at DESC, i.instance_id DESC
                LIMIT 100
                """
            )
            rows = await cur.fetchall()
        return [
            SmartReportInstanceRow(
                instance_id=int(r[0]),
                report_id=int(r[1]) if r[1] is not None else None,
                template_id=int(r[2]),
                template_name=str(r[3]) if r[3] is not None else None,
                instance_name=str(r[4]),
                generation_status=str(r[5]),
                output_file_path=str(r[6]) if r[6] is not None else None,
                error_message=str(r[7]) if r[7] is not None else None,
                last_generated_at=str(r[8]) if r[8] is not None else None,
                last_refresh_at=str(r[9]) if r[9] is not None else None,
                created_at=str(r[10]),
                updated_at=str(r[11]),
            )
            for r in rows
        ]

    async def instance_output_path(self, instance_id: int) -> Path:
        async with aiosqlite.connect(common_db_path()) as db:
            cur = await db.execute(
                "SELECT output_file_path FROM smart_report_instance WHERE instance_id = ?",
                (instance_id,),
            )
            row = await cur.fetchone()
        if not row or not row[0]:
            raise HTTPException(status_code=404, detail="报告实例没有可下载文件")
        path = Path(str(row[0]))
        if not path.exists() or not path.is_file():
            raise HTTPException(status_code=404, detail="报告文件不存在")
        return path

    def extract_placeholders(self, docx_path: Path) -> list[str]:
        doc = Document(str(docx_path))
        found: list[str] = []
        for text in self._iter_doc_text(doc):
            for match in PLACEHOLDER_RE.finditer(text):
                token = match.group(1).strip()
                if token and token not in found:
                    found.append(token)
        return found

    async def _sync_detected_variables(
        self, db: aiosqlite.Connection, template_id: int, placeholders: list[str], now: str
    ) -> None:
        for idx, token in enumerate(placeholders):
            variable_type, binding_config = self._infer_variable(token)
            await db.execute(
                """
                INSERT INTO smart_report_template_variable (
                  template_id, variable_key, variable_name, variable_type,
                  binding_config_json, display_order, created_at, updated_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(template_id, variable_key) DO UPDATE SET
                  display_order = excluded.display_order,
                  updated_at = excluded.updated_at
                """,
                (
                    template_id,
                    token,
                    self._default_variable_name(token),
                    variable_type,
                    json.dumps(binding_config, ensure_ascii=False),
                    idx,
                    now,
                    now,
                ),
            )

    def _infer_variable(self, token: str) -> tuple[str, dict[str, Any]]:
        text = token.strip()
        if ":" in text:
            prefix, raw_key = text.split(":", 1)
            prefix = prefix.strip().lower()
            key = raw_key.strip()
            if prefix == "metric":
                return "metric", {"metric_id": key}
            if prefix == "formula":
                parts = [p.strip() for p in key.split(":") if p.strip()]
                data_acct_code = parts[0] if parts else key
                formula_type = parts[1] if len(parts) > 1 else "budget"
                return "formula", {"data_acct_code": data_acct_code, "formula_type": formula_type}
            if prefix in {"calc", "calculated", "calculated_metric"}:
                return "calc", {"metric_code": key}
            if prefix in {"param", "parameter"}:
                return "parameter", {"param_key": key}
            if prefix in {"table", "chart", "analysis"}:
                return prefix, {f"{prefix}_id": key}
            if prefix == "text":
                return "text", {"text_key": key}
        return "text", {"text_key": text}

    def _default_variable_name(self, token: str) -> str:
        if ":" in token:
            prefix, key = token.split(":", 1)
            prefix_name = {
                "metric": "指标",
                "formula": "公式",
                "calc": "计算指标",
                "calculated": "计算指标",
                "calculated_metric": "计算指标",
                "param": "参数",
                "parameter": "参数",
                "table": "表格",
                "chart": "图表",
                "analysis": "分析结论",
                "text": "文本",
            }.get(prefix.strip().lower(), "变量")
            return f"{prefix_name} {key.strip()}"
        return token.strip()

    async def _fetch_template_record(self, template_id: int) -> dict[str, Any]:
        async with aiosqlite.connect(common_db_path()) as db:
            await db.execute("PRAGMA foreign_keys = ON")
            cur = await db.execute(
                "SELECT template_id, template_name, file_path FROM smart_report_template WHERE template_id = ?",
                (template_id,),
            )
            row = await cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="报告模板不存在")
        return {"template_id": int(row[0]), "template_name": str(row[1]), "file_path": str(row[2])}

    async def _resolve_values(
        self,
        variables: list[SmartReportTemplateVariableRow],
        parameters: dict[str, Any],
        text_values: dict[str, Any],
    ) -> tuple[dict[str, str], list[str]]:
        resolved: dict[str, str] = {}
        warnings: list[str] = []
        for variable in variables:
            token = variable.variable_key
            inferred_type, inferred_config = self._infer_variable(token)
            explicit_prefix = token.split(":", 1)[0].strip().lower() if ":" in token else ""
            variable_type = inferred_type if explicit_prefix in {"metric", "formula", "calc", "calculated", "calculated_metric", "param", "parameter", "text", "table", "chart", "analysis"} else variable.variable_type
            cfg = {**inferred_config, **(variable.binding_config or {})}
            if variable_type == "parameter":
                key = str(cfg.get("param_key") or token.split(":", 1)[-1]).strip()
                resolved[token] = str(parameters.get(key, ""))
            elif variable_type == "metric":
                metric_id = str(cfg.get("metric_id") or token.split(":", 1)[-1]).strip()
                value = await self._calculate_metric(metric_id, parameters)
                resolved[token] = value
            elif variable_type == "formula":
                resolved[token] = await self._resolve_formula_text(token, cfg, parameters)
            elif variable_type == "calc":
                metric_code = str(cfg.get("metric_code") or token.split(":", 1)[-1]).strip()
                resolved[token] = await self._calculate_report_metric(metric_code, parameters)
            elif variable_type == "text":
                key = str(cfg.get("text_key") or token.split(":", 1)[-1]).strip()
                resolved[token] = str(text_values.get(key, parameters.get(key, "")))
            else:
                resolved[token] = f"[{variable.variable_name} 待接入]"
                warnings.append(f"{variable.variable_key} 当前为占位输出，后续接入 {variable_type} 渲染")
        return resolved, warnings

    async def _resolve_formula_text(self, token: str, cfg: dict[str, Any], params: dict[str, Any]) -> str:
        raw = token.split(":", 1)[-1] if ":" in token else token
        parts = [p.strip() for p in raw.split(":") if p.strip()]
        data_acct_code = str(cfg.get("data_acct_code") or (parts[0] if parts else "")).strip().upper()
        formula_type = str(cfg.get("formula_type") or (parts[1] if len(parts) > 1 else "auto")).strip().lower()
        if formula_type == "auto":
            formula_type = "actual" if self._budget_actual_param(params) == 1 else "budget"
        if not data_acct_code:
            return ""
        async with aiosqlite.connect(common_db_path()) as db:
            cur = await db.execute(
                """
                SELECT data_acct_code, data_acct_name, budget_formula, actual_formula
                FROM data_account
                WHERE data_acct_code = ?
                """,
                (data_acct_code,),
            )
            row = await cur.fetchone()
        if not row:
            return f"{data_acct_code} 未找到公式"
        formula = row[3] if formula_type in {"actual", "1", "实际"} else row[2]
        label = "实际公式" if formula_type in {"actual", "1", "实际"} else "预算公式"
        return f"{row[0]} {row[1]} {label}：{formula or '未配置'}"

    async def _calculate_report_metric(self, metric_code: str, params: dict[str, Any]) -> str:
        metric = await self.get_calc_metric(metric_code)
        if not metric:
            return f"{metric_code} 未找到计算指标"
        values: dict[str, float] = {}
        for component in metric.components:
            values[component.alias] = await self._sum_data_account(component.data_acct_code, params, self._budget_actual_param(params))
        try:
            result = self._safe_eval_expression(metric.expression, values)
        except ZeroDivisionError:
            return "N/A"
        percentage = metric.format_type == "percent" or metric.value_type == "百分比"
        return _format_number(result, percentage=percentage)

    async def _calculate_metric(self, metric_id: str, params: dict[str, Any]) -> str:
        if metric_id == "m_budget_actual_gap":
            actual = await self._sum_budget_summary(params, budget_actual=1)
            budget = await self._sum_budget_summary(params, budget_actual=0)
            return _format_number(actual - budget)
        if metric_id == "m_yoy_growth":
            current = await self._sum_budget_summary(params, budget_actual=self._budget_actual_param(params))
            previous_params = dict(params)
            previous_params["year"] = _plain_year(params.get("year")) - 1
            previous = await self._sum_budget_summary(previous_params, budget_actual=self._budget_actual_param(params))
            if previous == 0:
                return "N/A"
            return _format_number((current - previous) / previous, percentage=True)
        if metric_id == "m_mom_growth":
            current_month = _month_label(params.get("month"))
            if not current_month:
                return "N/A"
            month_num = int(current_month[1:])
            if month_num <= 1:
                return "N/A"
            current = await self._sum_budget_summary(params, budget_actual=self._budget_actual_param(params))
            previous_params = dict(params)
            previous_params["month"] = f"M{month_num - 1:02d}"
            previous = await self._sum_budget_summary(previous_params, budget_actual=self._budget_actual_param(params))
            if previous == 0:
                return "N/A"
            return _format_number((current - previous) / previous, percentage=True)
        return "N/A"

    async def _sum_data_account(self, data_acct_code: str, params: dict[str, Any], budget_actual: int) -> float:
        scoped_params = dict(params)
        scoped_params["data_acct_code"] = data_acct_code
        return await self._sum_budget_summary(scoped_params, budget_actual=budget_actual)

    def _budget_actual_param(self, params: dict[str, Any]) -> int:
        raw = str(params.get("budget_actual", params.get("caliber", "1"))).strip().lower()
        if raw in {"0", "budget", "预算"}:
            return 0
        return 1

    async def _sum_budget_summary(self, params: dict[str, Any], *, budget_actual: int) -> float:
        year_int = _plain_year(params.get("year"))
        path = budget_db_path(year_int)
        if not path.exists():
            return 0.0
        where = ["year = ?", "budget_actual = ?"]
        values: list[Any] = [_year_label(year_int), budget_actual]
        month = _month_label(params.get("month"))
        start_month = _month_label(params.get("start_month"))
        end_month = _month_label(params.get("end_month"))
        if start_month or end_month:
            start_num = int((start_month or "M01")[1:])
            end_num = int((end_month or "M12")[1:])
            if start_num > end_num:
                start_num, end_num = end_num, start_num
            where.append("CAST(SUBSTR(month, 2) AS INTEGER) BETWEEN ? AND ?")
            values.extend([start_num, end_num])
        elif month:
            where.append("month = ?")
            values.append(month)
        quarter = str(params.get("quarter") or "").strip().upper()
        if quarter:
            where.append("quarter = ?")
            values.append(quarter)
        version_id = params.get("version_id")
        if version_id is not None and str(version_id).strip() != "":
            where.append("version_id = ?")
            values.append(int(version_id))
        data_code = params.get("data_acct_code") or params.get("data_account")
        if data_code:
            where.append("data_code_name LIKE ?")
            values.append(f"%{str(data_code).strip()}%")
        product_code = params.get("product_code") or params.get("product")
        if product_code:
            where.append("IFNULL(product_code_name, '') LIKE ?")
            values.append(f"%{str(product_code).strip()}%")
        dept = params.get("dept_code") or params.get("dept")
        if dept:
            where.append("(IFNULL(dept_level1, '') || IFNULL(dept_level2, '') || IFNULL(dept_level3, '')) LIKE ?")
            values.append(f"%{str(dept).strip()}%")
        report = params.get("report_code") or params.get("report")
        if report:
            where.append(
                "(IFNULL(report_level1, '') || IFNULL(report_level2, '') || IFNULL(report_level3, '') || IFNULL(report_level4, '') || IFNULL(report_level5, '')) LIKE ?"
            )
            values.append(f"%{str(report).strip()}%")

        async with aiosqlite.connect(path) as db:
            cur = await db.execute(
                f"SELECT COALESCE(SUM(value), 0) FROM budget_summary WHERE {' AND '.join(where)}",
                values,
            )
            row = await cur.fetchone()
        return float(row[0] or 0.0) if row else 0.0

    async def _variables_for_template(self, template_id: int, template_path: Path) -> list[SmartReportTemplateVariableRow]:
        rows = await self.list_variables(template_id)
        by_key = {item.variable_key: item for item in rows}
        now = _iso_now()
        for idx, token in enumerate(self.extract_placeholders(template_path)):
            if token in by_key:
                continue
            variable_type, binding_config = self._infer_variable(token)
            rows.append(
                SmartReportTemplateVariableRow(
                    variable_id=-(idx + 1),
                    template_id=template_id,
                    variable_key=token,
                    variable_name=self._default_variable_name(token),
                    variable_type=variable_type,  # type: ignore[arg-type]
                    binding_config=binding_config,
                    display_order=idx,
                    created_at=now,
                    updated_at=now,
                )
            )
        return rows

    def _calc_metric_from_row(self, row: Any) -> SmartReportCalcMetricRow:
        raw_components = json.loads(str(row[3] or "[]"))
        components = [
            SmartReportCalcMetricComponent(
                alias=str(item.get("alias") or ""),
                data_acct_code=str(item.get("data_acct_code") or ""),
                data_acct_name=str(item.get("data_acct_name")) if item.get("data_acct_name") is not None else None,
            )
            for item in raw_components
            if isinstance(item, dict)
        ]
        return SmartReportCalcMetricRow(
            metric_code=str(row[0]),
            metric_name=str(row[1]),
            expression=str(row[2]),
            components=components,
            value_type=str(row[4] or "金额"),
            format_type=str(row[5] or "number"),
            remark=str(row[6]) if row[6] is not None else None,
            created_at=str(row[7]),
            updated_at=str(row[8]),
        )

    def _safe_eval_expression(self, expression: str, values: dict[str, float]) -> float:
        node = ast.parse(expression, mode="eval")

        def visit(current: ast.AST) -> float:
            if isinstance(current, ast.Expression):
                return visit(current.body)
            if isinstance(current, ast.Constant) and isinstance(current.value, (int, float)):
                return float(current.value)
            if isinstance(current, ast.Name):
                if current.id not in values:
                    raise HTTPException(status_code=400, detail=f"表达式引用了未选择的基础指标：{current.id}")
                return float(values[current.id])
            if isinstance(current, ast.BinOp) and type(current.op) in _CALC_OPERATORS:
                return float(_CALC_OPERATORS[type(current.op)](visit(current.left), visit(current.right)))
            if isinstance(current, ast.UnaryOp) and type(current.op) in _CALC_OPERATORS:
                return float(_CALC_OPERATORS[type(current.op)](visit(current.operand)))
            raise HTTPException(status_code=400, detail="表达式只支持基础指标别名、数字和 + - * / ( ) 运算")

        return visit(node)

    def _render_docx(self, template_path: Path, output_path: Path, resolved: dict[str, str]) -> None:
        if not template_path.exists():
            raise HTTPException(status_code=404, detail="模板文件不存在")
        doc = Document(str(template_path))
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, resolved)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, resolved)
        doc.save(str(output_path))

    def _render_preview_text(self, template_path: Path, resolved: dict[str, str]) -> str:
        if not template_path.exists():
            raise HTTPException(status_code=404, detail="模板文件不存在")
        doc = Document(str(template_path))
        blocks: list[str] = []
        for paragraph in doc.paragraphs:
            text = self._replace_text(paragraph.text, resolved).strip()
            if text:
                blocks.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [self._replace_text(cell.text, resolved).strip() for cell in row.cells]
                line = " | ".join(cell for cell in cells if cell)
                if line:
                    blocks.append(line)
        return "\n".join(blocks)

    def _write_text_template_docx(self, target: Path, title: str, content: str) -> None:
        doc = Document()
        if title.strip():
            doc.add_heading(title.strip(), level=1)
        for line in content.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
            doc.add_paragraph(line)
        doc.save(str(target))

    def _replace_in_paragraph(self, paragraph: Any, resolved: dict[str, str]) -> None:
        text = paragraph.text
        if "{{" not in text:
            return
        new_text = self._replace_text(text, resolved)
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""

    def _replace_text(self, text: str, resolved: dict[str, str]) -> str:
        return PLACEHOLDER_RE.sub(lambda m: resolved.get(m.group(1).strip(), m.group(0)), text)

    def _iter_doc_text(self, doc: Any) -> list[str]:
        texts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.extend(p.text for p in cell.paragraphs)
        return texts
